# src/codecollate/collator.py

import logging
import re
import textwrap
import yaml
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import os
from dotenv import load_dotenv
from openai import OpenAI
from datetime import datetime

try:
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    docx = None

logger = logging.getLogger(__name__)

# --- DOCX 生成的核心目标参数 ---
TOTAL_TARGET_DOCX_PAGES = 63
ESTIMATED_LINES_PER_PAGE_FOR_DOCX = 57

DEFAULT_CONFIG = {
    'source_processing': {
        'include_extensions': ['*.html', '*.js', '*.css'],
        'strip_comments': True,
        'strip_blank_lines': True,
        'encoding_fallbacks': ['utf-8', 'utf-8-sig', 'gbk', 'latin-1'],
    },
    'expansion': {
        'enabled': True,
        'method': 'llm',  # llm | repeat | none
        'max_attempts': 5,
        'target_page_count': 100,
        'estimated_lines_per_page': 54,
        'logical_to_physical_ratio': 1.5,
        'safety_multiplier': 1.25,
        'repeat_marker': '/* === repeated block {index} === */',
        'fallback_to_repeat': True,
    },
    'wrapping': {
        'width': 85,
        'subsequent_indent': '  ',
    },
    'targets': {
        'docx_total_pages': 60,
        'docx_lines_per_page': 57,
    },
    # --- DOCX 相关的所有默认配置 ---
    'docx': {
        'page_setup': {
            'paper_size': 'A4', 
            'orientation': 'portrait', 
            'margins': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5}, 
            'header_footer_distance': {'header_from_top': 1.5, 'footer_from_bottom': 1.75}
        },
        'styles': {
            'code_block': {'font_name': 'Courier New', 'font_size': 9, 'line_spacing': 1.0, 'space_before': 0, 'space_after': 0}, 
            'header': {'font_name': 'Calibri', 'font_size': 10, 'alignment': 'CENTER'}, 
            'footer': {'font_name': 'Calibri', 'font_size': 10, 'alignment': 'CENTER'}
        },
        'header': {
            'content': '{software_name} {version}', 
            'border': {'enable': True, 'position': 'bottom', 'style': 'single', 'size': 4, 'space': 1, 'color': 'auto'}
        },
        'footer': {
            'content': '{page_number}'
        }
    },
}

class SourceCodeCollator:
    """
    一个专业的服务类，可将源代码高质量地转换为DOCX文档。
    完全由YAML配置文件驱动。
    """
    def __init__(self, source_dir: Path, software_name: str, version: str, output_dir: Path, config_path: str = None):
        if docx is None:
            raise ImportError("Missing dependencies. Please run 'uv pip install python-docx pyyaml openai \"python-dotenv[cli]\"'.")
        if not source_dir.is_dir():
            raise FileNotFoundError(f"Source directory not found: {source_dir}")

        self.source_dir = source_dir
        self.software_name = software_name
        self.version = version
        
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Output will be saved to: {self.output_dir.resolve()}")

        load_dotenv()
        self.config = self._load_config(config_path)
        self.llm_client = self._init_llm_client()

    def _init_llm_client(self) -> Optional[OpenAI]:
        """初始化 LLM 客户端（根据配置决定是否启用）。"""
        expansion_cfg = self.config.get('expansion', DEFAULT_CONFIG['expansion'])
        if not expansion_cfg.get('enabled', True):
            return None
        if expansion_cfg.get('method', 'llm') != 'llm':
            return None

        api_key = os.getenv("DASHSCOPE_API_KEY")
        if not api_key:
            logger.warning("DASHSCOPE_API_KEY not found; falling back to non-LLM expansion.")
            return None

        return OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        )
    def _load_config(self, config_path_str: str) -> Dict:
        """加载YAML配置，如果失败则使用默认配置。"""
        env_config_path = os.getenv("CODECOLLATE_CONFIG_PATH")
        if env_config_path:
            config_path = Path(env_config_path)
            if config_path.exists():
                logger.info(f"正在从环境变量 CODECOLLATE_CONFIG_PATH='{config_path}' 加载配置...")
                with open(config_path, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f)
            logger.warning(f"环境变量 CODECOLLATE_CONFIG_PATH 指向的文件不存在: {config_path}")

        if config_path_str:
            config_path = Path(config_path_str)
            if config_path.exists():
                logger.info(f"正在从 '{config_path}' 加载配置...")
                with open(config_path, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f)
        
        default_config_path = Path.cwd() / "codecollate_config.yaml"
        if default_config_path.exists():
            logger.info(f"正在从默认路径 '{default_config_path}' 加载配置...")
            with open(default_config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)

        logger.warning("未找到任何YAML配置文件，将使用内置的默认样式。")
        return DEFAULT_CONFIG

    def run(self) -> Path:
        """
        执行完整的文档生成流程。
        - DOCX采用“预处理物理行+精确估算”策略，让Word自然分页。
        """
        logger.info("Starting source code collation process...")
        
        clean_lines, _ = self._collect_and_clean_sources()
        if not clean_lines:
            raise ValueError("No processable source code found.")
        
        expanded_lines = self._expand_code_with_llm(clean_lines)

        # 关键步骤：将所有逻辑行预处理为适合页面宽度的物理行。
        # 这是确保我们对DOCX页数估算准确性的核心。
        logger.info("Preprocessing all code lines to wrap long lines for document generation...")
        physical_lines = []
        wrap_cfg = self.config.get('wrapping', DEFAULT_CONFIG['wrapping'])
        WRAP_LIMIT = wrap_cfg.get('width', 85)  # 一个相对保守的行宽，避免Word进行二次强制换行
        subsequent_indent = wrap_cfg.get('subsequent_indent', '  ')
        for line in expanded_lines:
            wrapped_lines = textwrap.wrap(
                line, 
                width=WRAP_LIMIT, 
                break_long_words=True, 
                break_on_hyphens=False, 
                subsequent_indent=subsequent_indent
            )
            physical_lines.extend(wrapped_lines if wrapped_lines else [''])
        
        total_physical_lines = len(physical_lines)
        
        desired_base_filename = f"{self.software_name}_{self.version}_源代码整理文档"
        base_filename = re.sub(r'[\\/*?:"<>|]', '_', desired_base_filename)

        # --- DOCX 流程 (最终的、基于物理行估算的稳定版本) ---
        docx_path = self.output_dir / f"{base_filename}.docx"
        
        # 1. 使用预处理过的 physical_lines 进行切片
        docx_lines = self._slice_lines_for_docx(physical_lines)
        
        # 2. 使用切片后的代码生成 DOCX，让Word自然分页
        self._generate_docx_from_config(docx_lines, docx_path)

        logger.info("Source code collation process completed successfully!")
        return docx_path

    def _collect_and_clean_sources(self) -> Tuple[List[str], Dict[str, int]]:
        """
        根据配置文件搜集、排序并处理所有源文件，智能识别文件类型并分别处理。
        """
        # --- 新增逻辑：从配置中读取要包含的文件扩展名 ---
        source_cfg = self.config.get('source_processing', DEFAULT_CONFIG['source_processing'])
        extensions = source_cfg.get('include_extensions', ['*.html', '*.js', '*.css'])
        strip_comments = source_cfg.get('strip_comments', True)
        strip_blank_lines = source_cfg.get('strip_blank_lines', True)
        encodings = source_cfg.get('encoding_fallbacks', ['utf-8', 'utf-8-sig', 'gbk', 'latin-1'])
        exclude_dirs = set(source_cfg.get('exclude_dirs', []))
        
        if not extensions:
            logger.warning("Configuration for 'include_extensions' is empty. No files will be processed.")
            return [], {}
            
        logger.info(f"Step 1: Collecting files with extensions: {', '.join(extensions)}")
        
        source_files = []
        for ext in extensions:
            source_files.extend(self.source_dir.rglob(ext))
        if exclude_dirs:
            source_files = [p for p in source_files if not self._is_excluded_path(p, exclude_dirs)]
        source_files.sort()
        all_clean_lines: List[str] = []
        source_files_info: Dict[str, int] = {}
        for file_path in source_files:
            try:
                code = self._read_text_with_fallback(file_path, encodings)
                if strip_comments:
                    code = self._strip_comments(code, file_path.suffix.lower())
                original_lines = code.splitlines()

                is_minified = len(original_lines) <= 10 and len(code) > 2000

                if is_minified:
                    logger.info(f"Detected minified file '{file_path.name}', applying special formatting...")
                    formatted_code = code.replace(';', ';\n')
                    lines = formatted_code.splitlines()
                    clean_lines = [line.strip() for line in lines if (not strip_blank_lines or line.strip())]
                else:
                    if strip_blank_lines:
                        clean_lines = [line for line in original_lines if line.strip()]
                    else:
                        clean_lines = list(original_lines)
                
                if clean_lines:
                    relative_path = str(file_path.relative_to(self.source_dir))
                    source_files_info[relative_path] = len(clean_lines)
                    all_clean_lines.extend(clean_lines)

            except Exception as e:
                logger.warning(f"Could not read or process file {file_path}: {e}")
        return all_clean_lines, source_files_info

    def _is_excluded_path(self, file_path: Path, exclude_dirs: set[str]) -> bool:
        """判断路径是否位于排除目录中。"""
        try:
            relative_parts = file_path.relative_to(self.source_dir).parts
        except ValueError:
            relative_parts = file_path.parts
        return any(part in exclude_dirs for part in relative_parts)

    def _read_text_with_fallback(self, file_path: Path, encodings: List[str]) -> str:
        """按顺序尝试多种编码读取文本，最后兜底替换无法解码的字符。"""
        for enc in encodings:
            try:
                return file_path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        logger.warning(f"Failed to decode {file_path} with common encodings; using replacement characters.")
        return file_path.read_text(encoding=encodings[0], errors='replace')

    def _strip_comments(self, code: str, suffix: str) -> str:
        """按文件类型做轻量级注释清理（不做完整语法解析）。"""
        if suffix in {'.html', '.htm'}:
            return re.sub(r'<!--.*?-->', '', code, flags=re.S)
        if suffix in {'.js', '.ts', '.java', '.c', '.cpp', '.cs', '.css'}:
            code = re.sub(r'/\*.*?\*/', '', code, flags=re.S)
            return re.sub(r'//.*', '', code)
        if suffix == '.py':
            lines = []
            for line in code.splitlines():
                if line.lstrip().startswith('#'):
                    continue
                lines.append(line)
            return "\n".join(lines)
        return code


    def _expand_code_with_llm(self, original_lines: List[str]) -> List[str]:
        """使用 LLM 扩充代码，必要时回退到重复扩充。"""
        expansion_cfg = self.config.get('expansion', DEFAULT_CONFIG['expansion'])
        if not expansion_cfg.get('enabled', True) or expansion_cfg.get('method', 'llm') == 'none':
            return original_lines

        target_page_count = expansion_cfg.get('target_page_count', 100)
        est_lines_per_page = expansion_cfg.get('estimated_lines_per_page', 54)
        logical_to_physical_ratio = expansion_cfg.get('logical_to_physical_ratio', 1.5)
        safety_multiplier = expansion_cfg.get('safety_multiplier', 1.25)

        estimated_physical_lines = target_page_count * est_lines_per_page
        estimated_logical_lines = int(estimated_physical_lines / logical_to_physical_ratio)
        safe_target_line_count = int(estimated_logical_lines * safety_multiplier)

        logger.info(
            f"Step 2: Expanding code to generate a full document of ~{target_page_count} pages "
            f"(target ~{safe_target_line_count} logical lines)..."
        )

        if not self.llm_client:
            return self._expand_code_by_repetition(original_lines, safe_target_line_count)

        current_lines = list(original_lines)  # 创建一个可变副本
        context_code = "\n".join(original_lines[:1000])
        max_attempts = expansion_cfg.get('max_attempts', 5)

        attempt = 0
        while len(current_lines) < safe_target_line_count and attempt < max_attempts:
            attempt += 1
            lines_to_generate = safe_target_line_count - len(current_lines)
            logger.info(
                f"--- Attempt {attempt}/{max_attempts}: Code still insufficient. "
                f"Needing {lines_to_generate} more lines. Calling LLM... ---"
            )

            system_prompt = """
You are an expert polyglot programmer, proficient in multiple programming languages.
Your task is to expand a codebase for documentation purposes by generating high-quality, realistic source code.
You must follow all constraints precisely. You must generate a large volume of code.
"""
            user_prompt = f"""
            **Context:**
            The existing code is a mix of web technologies. Your task is to generate new, self-contained code modules that would complement a project of this nature.
            
            **Constraints:**
            1. **Primary Goal - Quantity:** Your most important task is to generate a large volume of code. You MUST generate **at least {lines_to_generate}** new lines of source code. Do not stop early.
            2. **Language Choice:** You can generate code in any of the following languages: Python, Java, C, C++, C#, PHP. Choose a language and stick to it for the generated block.
            3. **Code Content:** The generated code must be different from previous generations and the provided context. Create new, unique functionality.
            4. **Style:** The code must be well-formatted, include comments where appropriate, and look professional.
            5. **CRITICAL: Only output the raw, new source code.** Do not include any explanations or markdown formatting.

            **Example of Existing Code (for style reference only):**
            ---
            {context_code}
            ---
            **New Source Code (in Python, Java, C, C++, C#, or PHP):**
            """

            try:
                response = self.llm_client.chat.completions.create(
                    model="qwen3-coder-flash",
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                    temperature=0.75,
                )
                generated_code = response.choices[0].message.content
                if generated_code:
                    generated_lines = generated_code.splitlines()
                    logger.info(f"Successfully generated {len(generated_lines)} new lines from Qwen on attempt {attempt}.")
                    current_lines.extend(generated_lines)
                    logger.info(f"Total lines are now: {len(current_lines)} / {safe_target_line_count}")
                else:
                    logger.warning(f"LLM returned an empty response on attempt {attempt}.")
                    break
            except Exception as e:
                logger.error(f"An exception occurred during LLM call on attempt {attempt}: {e}. Stopping expansion.")
                break

        if len(current_lines) < safe_target_line_count:
            logger.warning(
                f"LLM expansion finished after {attempt} attempts, but target of "
                f"{safe_target_line_count} lines was not reached."
            )
            if expansion_cfg.get('fallback_to_repeat', True):
                return self._expand_code_by_repetition(current_lines, safe_target_line_count)
        else:
            logger.info("LLM expansion successful. Total code quantity is now sufficient.")

        return current_lines

    def _expand_code_by_repetition(self, original_lines: List[str], target_line_count: int) -> List[str]:
        """用重复 + 标记的方式扩充代码行数，确保文档页数要求。"""
        expansion_cfg = self.config.get('expansion', DEFAULT_CONFIG['expansion'])
        marker_template = expansion_cfg.get('repeat_marker', '/* === repeated block {index} === */')

        if len(original_lines) >= target_line_count:
            return original_lines

        logger.info(
            f"Using repeat expansion to reach target of {target_line_count} lines "
            f"(current: {len(original_lines)})."
        )

        expanded = list(original_lines)
        index = 1
        while len(expanded) < target_line_count:
            expanded.append(marker_template.format(index=index))
            expanded.extend(original_lines)
            index += 1

        return expanded[:target_line_count]

    def _slice_lines_for_docx(self, all_physical_lines: List[str]) -> List[str]:
        """
        为DOCX文档精确准备源代码“物理行”，以生成一个约60页的文档。
        此方法基于校准后的每页物理行数进行估算。
        """
        # 1. 计算生成60页文档所需要的总物理行数
        targets_cfg = self.config.get('targets', DEFAULT_CONFIG['targets'])
        target_total_pages = targets_cfg.get('docx_total_pages', TOTAL_TARGET_DOCX_PAGES)
        estimated_lines_per_page = targets_cfg.get('docx_lines_per_page', ESTIMATED_LINES_PER_PAGE_FOR_DOCX)
        target_total_lines = target_total_pages * estimated_lines_per_page
        
        # 2. 计算前后两个部分各自需要多少行代码
        lines_per_part = target_total_lines // 2
        
        total_available_lines = len(all_physical_lines)

        logger.info(
            f"Slicing code for DOCX: aiming for {target_total_pages} pages "
            f"by selecting a total of {target_total_lines} physical lines "
            f"({lines_per_part} from the start and {lines_per_part} from the end)."
        )

        # 3. 检查总代码行数是否足够
        if total_available_lines < target_total_lines:
            logger.warning(
                f"Total available lines ({total_available_lines}) is less than the required "
                f"{target_total_lines} for a full 60-page document. "
                "The generated DOCX will contain all available code and may be shorter."
            )
            return all_physical_lines
        
        # 4. 精确切片：获取前、后部分的代码
        part1_lines = all_physical_lines[:lines_per_part]
        part2_lines = all_physical_lines[-lines_per_part:]
        
        # 5. 在中间添加一个清晰的分隔符
        separator = [
            "", "// " + "="*70, "// --- [ 内容跳跃：此处省略了文档中间部分源代码 ] ---",
            "// --- [ Content Omitted: Middle part of the source code is skipped here ] ---",
            "// " + "="*70, "",
        ]

        final_lines = part1_lines + separator + part2_lines
        logger.info(
            f"Sliced DOCX content prepared, total physical lines to be written: {len(final_lines)}."
        )
        return final_lines

    def _generate_docx_from_config(self, code_lines: List[str], docx_path: Path) -> None:
        """从零开始，完全根据加载的配置来生成DOCX文档。"""
        logger.info(f"Generating DOCX document based on configuration -> {docx_path}")
        
        cfg = self.config.get('docx', DEFAULT_CONFIG['docx'])
        document = docx.Document()
        
        section = document.sections[0]
        ps_cfg = cfg['page_setup']
        section.left_margin = Cm(ps_cfg['margins']['left'])
        section.right_margin = Cm(ps_cfg['margins']['right'])
        section.top_margin = Cm(ps_cfg['margins']['top'])
        section.bottom_margin = Cm(ps_cfg['margins']['bottom'])
        section.header_distance = Cm(ps_cfg['header_footer_distance']['header_from_top'])
        section.footer_distance = Cm(ps_cfg['header_footer_distance']['footer_from_bottom'])

        style = document.styles['Normal']
        cb_cfg = cfg['styles']['code_block']
        style.font.name = cb_cfg['font_name']
        style.font.size = Pt(cb_cfg['font_size'])
        style.paragraph_format.line_spacing = cb_cfg['line_spacing']
        style.paragraph_format.space_before = Pt(cb_cfg['space_before'])
        style.paragraph_format.space_after = Pt(cb_cfg['space_after'])

        header_cfg = cfg['header']
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = header_cfg['content'].format(software_name=self.software_name, version=self.version)
        header_p.alignment = WD_ALIGN_PARAGRAPH[cfg['styles']['header']['alignment']]
        if header_cfg.get('border', {}).get('enable'):
            p_pr = header_p._p.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            p_pr.append(p_bdr)
            bdr_cfg = header_cfg['border']
            bottom_bdr = OxmlElement(f'w:{bdr_cfg.get("position", "bottom")}')
            bottom_bdr.set(qn('w:val'), bdr_cfg.get('style', 'single'))
            bottom_bdr.set(qn('w:sz'), str(bdr_cfg.get('size', 4)))
            bottom_bdr.set(qn('w:space'), str(bdr_cfg.get('space', 1)))
            bottom_bdr.set(qn('w:color'), bdr_cfg.get('color', 'auto'))
            p_bdr.append(bottom_bdr)

        footer_cfg = cfg['footer']
        footer = section.footer
        footer_p = footer.paragraphs[0]
        if '{page_number}' in footer_cfg.get('content', ''):
             run = footer_p.add_run()
             fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin'); run._r.append(fldChar1)
             instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'; run._r.append(instrText)
             fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end'); run._r.append(fldChar2)
        footer_p.alignment = WD_ALIGN_PARAGRAPH[cfg['styles']['footer']['alignment']]

        for line in code_lines:
            document.add_paragraph(line)

        try:
            document.save(docx_path)
            logger.info("DOCX document generated successfully.")
        except Exception as e:
            logger.error(f"Failed to save DOCX file: {e}")
            raise RuntimeError("DOCX file generation failed") from e
